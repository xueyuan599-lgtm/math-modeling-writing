# -*- coding: utf-8 -*-
"""Markdown -> 国赛 Word (.docx) 转换器（md2word 技能核心）。

流程：预处理 md（剥离草稿痕迹、提取公式编号）-> pandoc 转 docx（LaTeX 公式转
Word 原生 OMML）-> python-docx/lxml 后处理（套用国赛样式、三线表、公式右编号、
图占位框、代码样式）。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

SKILL_ROOT = Path(__file__).resolve().parent.parent
LIB_DIR = SKILL_ROOT / "lib"
if str(LIB_DIR) not in sys.path:
    sys.path.insert(0, str(LIB_DIR))

try:
    import pypandoc
except ImportError:
    pypandoc = None

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

CHAPTER_NAMES = {
    "摘要",
    "问题重述",
    "问题分析",
    "模型假设",
    "符号说明",
    "模型建立与求解",
    "模型评价",
    "参考文献",
    "附录",
}

EDITORIAL_HEADING_RE = re.compile(
    r"(自检对照|自检清单|修改说明|可选升级|待确认|草稿说明|"
    r"v\s*\d+\s*[-—→>]|检查清单|对照表|写作过程|提示)"
)

DISPLAY_MATH_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
TAG_RE = re.compile(r"\\tag\s*\{?\s*(\d+)\s*\}?")
INLINE_MATH_RE = re.compile(r"\$([^$\n]+?)\$")


# ---------------------------------------------------------------------------
# Markdown 预处理
# ---------------------------------------------------------------------------

def normalize_heading(text: str) -> str:
    text = re.sub(r"^\s*[一二三四五六七八九十]+\s*[、.．]\s*", "", text.strip())
    return re.sub(r"\s+", "", text)


def preprocess_md(text: str, keep_editorial: bool) -> tuple[str, list[dict]]:
    """剥离草稿痕迹，返回 (处理后的 md, 已剔除清单)。"""
    removed: list[dict] = []
    if keep_editorial:
        return text, removed

    out: list[str] = []
    skip_until_level: int | None = None
    skipped_section: str | None = None

    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.lstrip()

        if skip_until_level is not None:
            m = re.match(r"^(#{1,6})\s+", stripped)
            if m and len(m.group(1)) <= skip_until_level:
                skip_until_level = None
                skipped_section = None
            else:
                continue

        if stripped.startswith(">"):
            removed.append({"type": "blockquote", "line": stripped[1:].strip()})
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            if EDITORIAL_HEADING_RE.search(heading_text):
                skip_until_level = level
                removed.append({"type": "section", "heading": heading_text})
                continue

        out.append(line)

    return "\n".join(out), removed


def extract_display_math(md_text: str, validate_tags: bool) -> tuple[str, list[dict]]:
    """剥离显示公式中的 \\tag{N}，记录编号与显示公式顺序索引。"""
    tags: list[dict] = []
    display_index = 0

    def repl(m: re.Match) -> str:
        nonlocal display_index
        src = m.group(1)
        tag_m = TAG_RE.search(src)
        num = None
        if tag_m:
            num = int(tag_m.group(1))
            src = src[: tag_m.start()] + src[tag_m.end() :]
        tags.append({"index": display_index, "number": num, "source": src.strip()})
        display_index += 1
        return "$$" + src + "$$"

    cleaned = DISPLAY_MATH_RE.sub(repl, md_text)

    if validate_tags:
        numbers = [t["number"] for t in tags if t["number"] is not None]
        expected = list(range(1, len(numbers) + 1))
        if numbers != expected:
            raise ValueError(
                "公式 \\tag 编号不连续："
                + "、".join(str(n) if n is not None else "无编号" for n in numbers)
            )
    return cleaned, tags


def extract_inline_math(md_text: str) -> list[str]:
    return [m.group(1).strip() for m in INLINE_MATH_RE.finditer(md_text)]


# ---------------------------------------------------------------------------
# Word 样式
# ---------------------------------------------------------------------------

def set_east_asia(style_or_run, east: str, latin: str | None = None) -> None:
    if latin:
        style_or_run.font.name = latin
    rpr = style_or_run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    if latin:
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)


def set_spacing_lines(par_or_style, before: int = 0, after: int = 0) -> None:
    pf = par_or_style.paragraph_format
    if before:
        pf.space_before = Pt(before * 12)
    if after:
        pf.space_after = Pt(after * 12)
    ppr = par_or_style._element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    if before:
        spacing.set(qn("w:beforeLines"), str(before * 50))
    if after:
        spacing.set(qn("w:afterLines"), str(after * 50))


def set_first_line_chars(par_or_style, chars: int) -> None:
    ppr = par_or_style._element.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if chars:
        ind.set(qn("w:firstLineChars"), str(chars * 100))
        ind.set(qn("w:firstLine"), str(chars * 240))
    else:
        ind.set(qn("w:firstLineChars"), "0")
        ind.set(qn("w:firstLine"), "0")


def add_paragraph_style(doc: Document, name: str, base: str | None = None) -> None:
    if name in {s.name for s in doc.styles}:
        return
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        try:
            st.base_style = doc.styles[base]
        except KeyError:
            pass
    st.quick_style = True


def define_guoshi_styles(doc: Document) -> None:
    """按 math-modeling-writing/references/typography.md 定义国赛样式。"""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    set_east_asia(normal, "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    specs = [
        ("题目", 16, "黑体", "Times New Roman", False, "center", 0.5, 0.5, 0),
        ("摘要标题", 14, "黑体", "Times New Roman", False, "center", 0.5, 0.5, 0),
        ("摘要正文", 12, "宋体", "Times New Roman", False, "justify", 0, 0, 2),
        ("正文", 12, "宋体", "Times New Roman", False, "justify", 0, 0, 2),
        ("一级标题", 14, "黑体", "Times New Roman", False, "center", 0.5, 0.5, 0),
        ("二级标题", 12, "黑体", "Times New Roman", False, "left", 0.5, 0.5, 0),
        ("三级标题", 12, "黑体", "Times New Roman", False, "left", 0.5, 0.5, 0),
        ("图注", 10.5, "宋体", "Times New Roman", True, "center", 0, 0.5, 0),
        ("表注", 10.5, "宋体", "Times New Roman", True, "center", 0.5, 0, 0),
        ("表格文字", 10.5, "宋体", "Times New Roman", False, "center", 0.5, 0, 0),
        ("列表文字", 12, "宋体", "Times New Roman", False, "left", 0, 0, 0),
        ("代码", 10.5, "宋体", "Times New Roman", False, "left", 0, 0, 0),
    ]
    for name, size, east, latin, bold, align, before, after, indent_chars in specs:
        add_paragraph_style(doc, name, None if name == "表格文字" else "Normal")
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = bold
        set_east_asia(st, east, latin)
        st.paragraph_format.line_spacing = 1.5
        align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        st.paragraph_format.alignment = align_map[align]
        set_spacing_lines(st, before, after)
        set_first_line_chars(st, indent_chars)

    code = doc.styles["代码"]
    code.paragraph_format.line_spacing = 1.5
    ppr = code.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    ppr.append(shd)

    tbl = doc.styles["表格文字"]
    tbl.paragraph_format.line_spacing = 1.5
    ppr = tbl.element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:beforeLines"), "50")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:afterLines"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")


# ---------------------------------------------------------------------------
# OOXML 后处理
# ---------------------------------------------------------------------------

def _insert_before_first(container, new_el, tag_names: list[str]) -> None:
    for tag in tag_names:
        found = container.find(qn(tag))
        if found is not None:
            found.addprevious(new_el)
            return
    container.append(new_el)


def _paragraph_style_name(p_el) -> str | None:
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return None
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        return None
    return pstyle.get(qn("w:val"))


def _paragraph_text(p_el) -> str:
    """取段落可见文本（含公式 m:t，按文档顺序）。"""
    parts: list[str] = []
    for node in p_el.iter():
        if node.tag == qn("w:t") or node.tag == qn("m:t"):
            parts.append(node.text or "")
    return "".join(parts)


def _set_paragraph_props(par, style_name: str | None = None, align=None,
                         indent_chars: int | None = None, doc: Document | None = None) -> None:
    if style_name:
        if doc is not None:
            par.style = doc.styles[style_name]
        else:
            par.style = par.part.document.styles[style_name]
    if align:
        par.alignment = align
    if indent_chars is not None:
        set_first_line_chars(par._element, indent_chars)


def _set_placeholder_box(par) -> None:
    """【图占位：…】→ 浅灰底纹居中占位框。"""
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ppr = par._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "BFBFBF")
        pbdr.append(el)
    ppr.append(pbdr)
    for run in par.runs:
        run.font.size = Pt(10.5)
        set_east_asia(run, "宋体", "Times New Roman")


def _text_width(s: str) -> int:
    return sum(2 if ord(ch) > 127 else 1 for ch in s)


def _set_three_line_table(tbl, doc: Document, usable_width_dxa: int) -> None:
    """表格 → 三线表（顶/底 1.5 磅、表头下线 1.0 磅）、单元格居中、表头加粗。"""
    tbl_pr = tbl._tbl.tblPr
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), "000000")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    for side in ("left", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + side)
        el.set(qn("w:val"), "none")
        borders.append(el)
    borders.append(top)
    borders.append(bottom)
    _insert_before_first(
        tbl_pr, borders,
        ["w:tblLook", "w:tblLayout", "w:tblCellMar", "w:tblInd", "w:jc", "w:tblW"],
    )

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        _insert_before_first(
            tbl_pr, tbl_w,
            ["w:jc", "w:tblCellSpacing", "w:tblInd", "w:tblBorders",
             "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"],
        )
    tbl_w.set(qn("w:w"), str(usable_width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        _insert_before_first(tbl_pr, layout, ["w:tblCellMar", "w:tblLook"])
    layout.set(qn("w:type"), "fixed")

    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        _insert_before_first(tbl_pr, jc, ["w:tblLook", "w:tblLayout", "w:tblCellMar"])
    jc.set(qn("w:val"), "center")

    rows = tbl.rows
    if not rows:
        return

    # 按内容宽度比例分配列宽；内容过宽时整表降到小五号（9pt）再算
    grid = tbl._tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol")) if grid is not None else []
    if cols:
        ncols = len(cols)
        font_size = 10.5
        unit = 105  # 10.5pt 下每个西文字符约 105 twips
        max_units = [1] * ncols
        for row in rows:
            cells = row.cells
            for i, cell in enumerate(cells[:ncols]):
                max_units[i] = max(max_units[i], _text_width(cell.text) + 3)

        def col_need(units: int) -> int:
            return units * unit + 216 + 60  # 内容宽 + 单元格边距 + 余量

        if sum(col_need(u) for u in max_units) > usable_width_dxa:
            font_size = 9.0
            unit = 90
        total_need = sum(col_need(u) for u in max_units)
        widths = [
            max(1, int(usable_width_dxa * col_need(u) / total_need))
            for u in max_units
        ]
        widths[-1] += usable_width_dxa - sum(widths)
        for col_el, width in zip(cols, widths):
            col_el.set(qn("w:w"), str(width))
        for row in rows:
            for i, cell in enumerate(row.cells[:ncols]):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    _insert_before_first(
                        tc_pr, tc_w,
                        ["w:tcBorders", "w:shd", "w:vMerge", "w:gridSpan"],
                    )
                tc_w.set(qn("w:w"), str(widths[i]))
                tc_w.set(qn("w:type"), "dxa")

        def apply_cell_font(par) -> None:
            for run in par.runs:
                run.font.size = Pt(font_size)

        header = rows[0]
        for cell in header.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            for old in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(old)
            tc_borders = OxmlElement("w:tcBorders")
            btm = OxmlElement("w:bottom")
            btm.set(qn("w:val"), "single")
            btm.set(qn("w:sz"), "8")
            btm.set(qn("w:space"), "0")
            btm.set(qn("w:color"), "000000")
            tc_borders.append(btm)
            _insert_before_first(tc_pr, tc_borders, ["w:shd", "w:vMerge", "w:tcW"])
            for p in cell.paragraphs:
                _set_paragraph_props(p, "表格文字", doc=doc)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_cell_font(p)
                for run in p.runs:
                    run.font.bold = True
        for row in rows[1:]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    _set_paragraph_props(p, "表格文字", doc=doc)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    apply_cell_font(p)
    else:
        header = rows[0]
        for cell in header.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                _set_paragraph_props(p, "表格文字", doc=doc)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
        for row in rows[1:]:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    _set_paragraph_props(p, "表格文字", doc=doc)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _number_display_math(doc: Document, tags: list[dict], usable_width_dxa: int) -> None:
    """显示公式段落：居中 + 右侧制表位 + 追加 (N)。"""
    body = doc.element.body
    tag_by_index = {t["index"]: t["number"] for t in tags if t["number"] is not None}
    math_paras: list = []
    for child in body.iterchildren():
        if child.tag == qn("w:p") and child.find(qn("m:oMathPara")) is not None:
            math_paras.append(child)

    max_index = max(tag_by_index, default=-1)
    if len(math_paras) <= max_index:
        raise RuntimeError(
            f"显示公式数量不匹配：md 中第 {max_index + 1} 个显示公式未生成 OMML"
            f"（docx 中仅 {len(math_paras)} 个）"
        )

    for i, p_el in enumerate(math_paras):
        num = tag_by_index.get(i)
        # 块级 oMathPara -> 行内 oMath（标准“公式居中+编号右对齐”段落结构）
        o_math_para = p_el.find(qn("m:oMathPara"))
        o_math = o_math_para.find(qn("m:oMath"))
        if o_math is not None:
            o_math_para.addprevious(o_math)
            p_el.remove(o_math_para)
        ppr = p_el.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            p_el.insert(0, ppr)
        jc = ppr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            ppr.append(jc)
        jc.set(qn("w:val"), "center")
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:firstLineChars"), "0")
        ind.set(qn("w:firstLine"), "0")
        tabs = ppr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            ppr.append(tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(usable_width_dxa))
        tabs.append(tab)
        if num is not None:
            run = OxmlElement("w:r")
            tab_run = OxmlElement("w:tab")
            t = OxmlElement("w:t")
            t.text = f"({num})"
            run.append(tab_run)
            run.append(t)
            p_el.append(run)


def _map_paragraph_styles(doc: Document, title: str | None,
                          usable_width_dxa: int) -> None:
    """按 md 章节约定映射 pandoc 段落到国赛样式。"""
    body = doc.element.body
    blocks: list = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            blocks.append(("p", child))
        elif child.tag == qn("w:tbl"):
            blocks.append(("tbl", child))

    from docx.text.paragraph import Paragraph
    from docx.table import Table

    paragraphs = [Paragraph(el, doc) for kind, el in blocks if kind == "p"]
    tables = [Table(el, doc) for kind, el in blocks if kind == "tbl"]

    if title:
        p = doc.add_paragraph(style=doc.styles["题目"])
        run = p.add_run(title)
        set_east_asia(run, "黑体", "Times New Roman")
        run.font.size = Pt(16)
        body.insert(0, p._p)

    # 标题角色判定
    heading_roles: dict[int, str] = {}
    first_h1 = True
    in_abstract = False
    in_references = False

    # 先记录“表前段落”的块索引（用于表注识别）
    table_prev_idx: set[int] = set()
    prev_kind = None
    for idx, (kind, el) in enumerate(blocks):
        if kind == "tbl":
            if prev_kind == "p" and idx > 0:
                table_prev_idx.add(idx - 1)
        prev_kind = kind

    tbl_counter = 0
    for idx, (kind, el) in enumerate(blocks):
        if kind == "tbl":
            _set_three_line_table(tables[tbl_counter], doc, usable_width_dxa)
            tbl_counter += 1
            continue
        par = Paragraph(el, doc)
        pstyle = _paragraph_style_name(el)
        text = par.text.strip()
        norm = normalize_heading(text)

        if pstyle and pstyle.startswith("Heading"):
            level = int(pstyle.split()[-1]) if pstyle.split()[-1].isdigit() else 1
            if norm == "摘要":
                heading_roles[idx] = "摘要标题"
                in_abstract = True
                in_references = False
            elif norm == "参考文献":
                heading_roles[idx] = "一级标题"
                in_abstract = False
                in_references = True
            elif level == 1:
                if first_h1 and norm not in CHAPTER_NAMES:
                    heading_roles[idx] = "题目"
                else:
                    heading_roles[idx] = "一级标题"
                first_h1 = False
                in_abstract = False
                in_references = False
            elif level == 2:
                heading_roles[idx] = "二级标题"
                in_abstract = False
                in_references = False
            else:
                heading_roles[idx] = "三级标题"
                in_abstract = False
                in_references = False
            continue

        # 非标题段落
        if pstyle and pstyle.replace(" ", "") == "SourceCode":
            heading_roles[idx] = "代码"
            continue
        if pstyle and pstyle.replace(" ", "") == "Caption":
            heading_roles[idx] = "图注"
            continue
        if text.startswith("【图占位"):
            heading_roles[idx] = "placeholder"
            continue
        if re.match(r"^图\s*\d+[：:]", text) or re.match(r"^图\s*\d+\s+", text):
            heading_roles[idx] = "图注"
            continue
        if idx in table_prev_idx and re.match(r"^表\s*\d+\s+", text):
            heading_roles[idx] = "表注"
            continue
        ppr = el.find(qn("w:pPr"))
        if ppr is not None and ppr.find(qn("w:numPr")) is not None:
            heading_roles[idx] = "列表文字"
            continue
        if in_abstract:
            heading_roles[idx] = "摘要正文"
        elif in_references:
            heading_roles[idx] = "正文-nolist"
        else:
            heading_roles[idx] = "正文"

    # 应用角色
    from docx.text.paragraph import Paragraph as P2

    for idx, (kind, el) in enumerate(blocks):
        if kind != "p":
            continue
        par = P2(el, doc)
        role = heading_roles.get(idx, "正文")
        if role == "placeholder":
            _set_paragraph_props(par, "正文", doc=doc)
            _set_placeholder_box(par)
        elif role == "正文-nolist":
            _set_paragraph_props(par, "正文", doc=doc)
            set_first_line_chars(par, 0)
        elif role == "代码":
            _set_paragraph_props(par, "代码", doc=doc)
        else:
            _set_paragraph_props(par, role, doc=doc)

        # 图片段落居中
        if el.find(".//" + qn("w:drawing")) is not None:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def convert(md_path: Path, out_path: Path, title: str | None = None,
            keep_editorial: bool = False, validate_tags: bool = True) -> dict:
    if pypandoc is None:
        raise RuntimeError(
            "缺少 pypandoc。请先执行：\n"
            f"  python -m pip install --target {LIB_DIR} pypandoc_binary"
        )

    md_text = md_path.read_text(encoding="utf-8")
    processed, removed = preprocess_md(md_text, keep_editorial)
    processed, tags = extract_display_math(processed, validate_tags)

    work = out_path.parent / f".md2word_{out_path.stem}"
    work.mkdir(parents=True, exist_ok=True)
    proc_md = work / "processed.md"
    raw_docx = work / "raw.docx"
    proc_md.write_text(processed, encoding="utf-8")

    extra_args = [
        "-f", "markdown-smart",
        "--wrap=none",
        f"--resource-path={md_path.parent}",
    ]
    pypandoc.convert_file(str(proc_md), "docx", outputfile=str(raw_docx),
                          extra_args=extra_args)

    doc = Document(str(raw_docx))
    define_guoshi_styles(doc)

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    usable_width_dxa = int(
        section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    )

    _map_paragraph_styles(doc, title, usable_width_dxa)
    _number_display_math(doc, tags, usable_width_dxa)

    doc.core_properties.title = title or md_path.stem
    doc.save(str(out_path))

    removed_path = out_path.with_name(out_path.stem + ".removed.json")
    removed_path.write_text(
        json.dumps({"removed": removed}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return {
        "output": str(out_path),
        "display_math": len(tags),
        "numbered_math": sum(1 for t in tags if t["number"] is not None),
        "removed": removed,
        "removed_report": str(removed_path),
    }


def main() -> int:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, ValueError):
            pass
    parser = argparse.ArgumentParser(description="Markdown -> 国赛 Word 转换器")
    parser.add_argument("input", type=Path, help="输入 md 文件")
    parser.add_argument("--out", type=Path, default=None, help="输出 docx 路径")
    parser.add_argument("--title", default=None, help="指定论文题目（覆盖首个 # 标题判定）")
    parser.add_argument("--keep-editorial", action="store_true",
                        help="保留自检对照/引用说明等草稿痕迹")
    parser.add_argument("--no-validate-tags", action="store_true",
                        help="跳过公式 \\tag 连续性校验")
    args = parser.parse_args()

    out = args.out or args.input.with_suffix(".docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        result = convert(
            args.input.resolve(),
            out.resolve(),
            title=args.title,
            keep_editorial=args.keep_editorial,
            validate_tags=not args.no_validate_tags,
        )
    except Exception as exc:  # noqa: BLE001
        print(f"转换失败：{exc}", file=sys.stderr)
        return 1

    print(json.dumps(result, ensure_ascii=False, indent=2))
    print(f"\n已生成：{out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
