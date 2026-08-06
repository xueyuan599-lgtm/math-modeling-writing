# -*- coding: utf-8 -*-
"""md2word 无损校验：内容逐类比对 + 结构审计 + 公式编号校验。"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import md2word
from md2word import (CHAPTER_NAMES, DISPLAY_MATH_RE, INLINE_MATH_RE,
                     normalize_heading, preprocess_md, extract_display_math)

from docx import Document
from docx.oxml.ns import qn


def latex_to_plain(s: str) -> str:
    greek = {
        "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
        "epsilon": "ε", "varepsilon": "ε", "zeta": "ζ", "eta": "η",
        "theta": "θ", "vartheta": "ϑ", "iota": "ι", "kappa": "κ",
        "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ", "pi": "π",
        "rho": "ρ", "sigma": "σ", "tau": "τ", "upsilon": "υ",
        "phi": "φ", "varphi": "φ", "chi": "χ", "psi": "ψ", "omega": "ω",
        "Gamma": "Γ", "Delta": "Δ", "Theta": "Θ", "Lambda": "Λ",
        "Xi": "Ξ", "Pi": "Π", "Sigma": "Σ", "Phi": "Φ", "Psi": "Ψ",
        "Omega": "Ω",
        "cdot": "·", "times": "×", "le": "≤", "ge": "≥", "neq": "≠",
        "approx": "≈", "in": "∈", "infty": "∞", "Rightarrow": "⇒",
        "rightarrow": "→", "leftarrow": "←",
    }
    for name, ch in greek.items():
        s = s.replace("\\" + name, ch)
    s = s.replace("\\{", "\ue000").replace("\\}", "\ue001")
    s = re.sub(r"\\(?:[a-zA-Z]+)", "", s)
    s = s.replace("{", "").replace("}", "").replace("_", "").replace("^", "")
    s = s.replace("\ue000", "{").replace("\ue001", "}")
    return s


def md_inline_to_text(s: str) -> str:
    s = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"\1", s)
    s = re.sub(r"\[([^\]]*)\]\([^)]+\)", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = s.replace("**", "").replace("__", "").replace("*", "").replace("_", "")
    s = INLINE_MATH_RE.sub(lambda m: latex_to_plain(m.group(1)), s)
    return s.strip()


def normalize(s: str) -> str:
    s = s.replace("−", "-").replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", "", s)


def extract_md_categories(md_text: str) -> dict:
    cats = {
        "headings": [],
        "paragraphs": [],
        "table_cells": [],
        "code_lines": [],
        "images": [],
        "placeholders": [],
        "captions": [],
        "display_math": 0,
        "inline_math": 0,
        "tables": 0,
    }
    in_code = False
    in_display_math = False
    code_buf: list[str] = []

    def flush_code() -> None:
        cats["code_lines"].extend(code_buf)
        code_buf.clear()

    for raw in md_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        is_math_line = "$$" in line
        if is_math_line:
            if not in_display_math:
                remaining = DISPLAY_MATH_RE.sub("", line)
                if remaining.strip():
                    cats["paragraphs"].append(md_inline_to_text(remaining))
            if line.count("$$") % 2 == 1:
                in_display_math = not in_display_math
            continue
        if not stripped:
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            cats["headings"].append(normalize_heading(m.group(2)))
            continue
        if re.fullmatch(r"\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*", stripped):
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            cats["table_cells"].extend(cells)
            continue
        img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if img:
            cats["images"].append(img.group(1).strip())
            alt = img.group(1).strip()
            if re.match(r"^(图|表)\s*\d+", alt):
                cats["captions"].append(alt)
            continue
        if stripped.startswith("【图占位"):
            cats["placeholders"].append(stripped)
            continue
        if re.match(r"^(图|表)\s*\d+", stripped):
            cats["captions"].append(stripped)
            continue
        if re.match(r"^\s*(?:[-*+]|\d+[.、])\s+", stripped):
            stripped = re.sub(r"^\s*(?:[-*+]|\d+[.、])\s+", "", stripped)
        cats["paragraphs"].append(md_inline_to_text(stripped))

    flush_code()
    cats["display_math"] = len(DISPLAY_MATH_RE.findall(md_text))
    cats["inline_math"] = len(INLINE_MATH_RE.findall(
        DISPLAY_MATH_RE.sub("", md_text)
    ))
    cats["tables"] = sum(
        1 for line in md_text.splitlines()
        if re.fullmatch(r"\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*", line.strip())
    )
    return cats


def para_text(p_el) -> str:
    parts: list[str] = []
    for node in p_el.iter():
        if node.tag == qn("w:t") or node.tag == qn("m:t"):
            parts.append(node.text or "")
        elif node.tag == qn("w:br"):
            parts.append("\n")
    return "".join(parts)


def is_display_math_para(p_el) -> bool:
    if p_el.find(qn("m:oMath")) is None:
        return False
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return False
    jc = ppr.find(qn("w:jc"))
    return jc is not None and jc.get(qn("w:val")) == "center"


def extract_docx_categories(doc: Document) -> dict:
    cats = {
        "headings": [],
        "paragraphs": [],
        "table_cells": [],
        "code_lines": [],
        "images": [],
        "placeholders": [],
        "captions": [],
        "display_math": 0,
        "inline_math": 0,
        "tables": len(doc.tables),
        "o_math_paras": [],
    }
    body = doc.element.body
    inline_math = 0
    for child in body.iterchildren():
        if child.tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                for tc in row.iter(qn("w:tc")):
                    texts = []
                    for p in tc.iter(qn("w:p")):
                        t = para_text(p).strip()
                        if t:
                            texts.append(t)
                    cats["table_cells"].append(" ".join(texts))
            continue
        if child.tag != qn("w:p"):
            continue
        pstyle = md2word._paragraph_style_name(child)
        text = para_text(child).strip()
        if is_display_math_para(child):
            cats["display_math"] += 1
            cats["o_math_paras"].append(text)
            continue
        inline_math += sum(
            1 for _ in child.iter(qn("m:oMath"))
        )
        if child.find(".//" + qn("w:drawing")) is not None:
            cats["images"].append("")
            if text:
                cats["captions"].append(text)
            continue
        if pstyle in ("题目", "摘要标题", "一级标题", "二级标题", "三级标题"):
            cats["headings"].append(normalize_heading(text))
        elif pstyle == "图注":
            cats["captions"].append(text)
        elif pstyle == "表注":
            cats["captions"].append(text)
        elif pstyle in ("代码", "SourceCode", "Source Code"):
            cats["code_lines"].extend(text.split("\n"))
        elif text.startswith("【图占位"):
            cats["placeholders"].append(text)
        elif text:
            cats["paragraphs"].append(text)
    cats["inline_math"] = inline_math
    return cats


def audit_docx(doc: Document) -> list[str]:
    errors: list[str] = []
    style_names = {s.name for s in doc.styles}
    body = doc.element.body
    tbl_idx = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            pstyle = md2word._paragraph_style_name(child)
            if pstyle and pstyle not in style_names:
                errors.append(f"未知段落样式：{pstyle}")
        elif child.tag == qn("w:tbl"):
            tbl = doc.tables[tbl_idx]
            tbl_idx += 1
            tbl_pr = tbl._tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                errors.append("表格缺少 tblBorders")
                continue
            top = borders.find(qn("w:top"))
            bottom = borders.find(qn("w:bottom"))
            if top is None or top.get(qn("w:sz")) != "12":
                errors.append("表格顶线不是 1.5 磅")
            if bottom is None or bottom.get(qn("w:sz")) != "12":
                errors.append("表格底线不是 1.5 磅")
            if tbl.rows:
                header_tc = tbl.rows[0].cells[0]._tc
                tc_pr = header_tc.find(qn("w:tcPr"))
                tc_borders = tc_pr.find(qn("w:tcBorders")) if tc_pr is not None else None
                if tc_borders is None:
                    errors.append("表头缺少 tcBorders（表头下线）")
                else:
                    btm = tc_borders.find(qn("w:bottom"))
                    if btm is None or btm.get(qn("w:sz")) != "8":
                        errors.append("表头下线不是 1.0 磅")
    return errors


def compare_cat(name: str, md_vals, docx_vals) -> list[str]:
    md_c = Counter(normalize(v) for v in md_vals)
    dx_c = Counter(normalize(v) for v in docx_vals)
    missing = list((md_c - dx_c).elements())
    extra = list((dx_c - md_c).elements())
    out = []
    if missing:
        out.append(f"{name} 缺失 {len(missing)} 项：{missing[:10]}")
    if extra:
        out.append(f"{name} 多出 {len(extra)} 项：{extra[:10]}")
    return out


def main() -> int:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, ValueError):
            pass
    parser = argparse.ArgumentParser(description="md2word 无损校验")
    parser.add_argument("input", type=Path, help="原始 md")
    parser.add_argument("out", type=Path, help="生成的 docx")
    parser.add_argument("--report", type=Path, default=None, help="JSON 报告路径")
    parser.add_argument("--keep-editorial", action="store_true")
    parser.add_argument("--no-validate-tags", action="store_true")
    args = parser.parse_args()

    errors: list[str] = []
    md_text = args.input.read_text(encoding="utf-8")
    processed, _ = preprocess_md(md_text, args.keep_editorial)
    _, tags = extract_display_math(processed, not args.no_validate_tags)
    md_cats = extract_md_categories(processed)
    doc = Document(str(args.out))
    dx_cats = extract_docx_categories(doc)

    # 内容比对
    for cat in ("headings", "paragraphs", "table_cells", "code_lines",
                "placeholders", "captions"):
        errors += compare_cat(cat, md_cats[cat], dx_cats[cat])

    # 图片：md 有 alt 文本，docx 统计数量
    if len(md_cats["images"]) != len(dx_cats["images"]):
        errors.append(
            f"图片数量不一致：md {len(md_cats['images'])}，"
            f"docx {len(dx_cats['images'])}"
        )

    # 表格数量
    if md_cats["tables"] != dx_cats["tables"]:
        errors.append(
            f"表格数量不一致：md {md_cats['tables']}，docx {dx_cats['tables']}"
        )

    # 公式数量
    if md_cats["display_math"] != dx_cats["display_math"]:
        errors.append(
            f"显示公式数量不一致：md {md_cats['display_math']}，"
            f"docx {dx_cats['display_math']}"
        )
    nums = [t["number"] for t in tags if t["number"] is not None]
    if nums and nums != list(range(1, len(nums) + 1)):
        errors.append(f"公式编号不连续：{nums}")

    # 结构审计
    errors += audit_docx(doc)

    report = {
        "ok": not errors,
        "errors": errors[:50],
        "md": {k: (len(v) if isinstance(v, list) else v) for k, v in md_cats.items()},
        "docx": {k: (len(v) if isinstance(v, list) else v) for k, v in dx_cats.items()},
        "tag_numbers": nums,
    }
    if args.report:
        args.report.write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if errors:
        print("\n校验失败：", file=sys.stderr)
        for e in errors[:50]:
            print(f"  - {e}", file=sys.stderr)
        return 1
    print("\n校验通过")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
